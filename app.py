import io
import re
from datetime import datetime
from typing import Any, Dict, List

import pandas as pd
import streamlit as st
from docx import Document

try:
    import fitz  # PyMuPDF
except Exception:
    fitz = None


APP_TITLE = "EDC智能稽查与RBQM风险识别系统"
st.set_page_config(page_title=APP_TITLE, layout="wide")
st.title(APP_TITLE)
st.caption("MVP：读取EDC Excel多Sheet、自动字段识别、基础数据核查、RBQM风险提示、Word报告导出。")


FIELD_SYNONYMS = {
    "subject_id": ["subject id", "subject no", "subjid", "受试者编号", "受试者id", "受试者", "筛选号", "随机号"],
    "site_id": ["site", "site no", "center", "中心号", "中心编号", "研究中心"],
    "visit_name": ["visit", "visit name", "folder", "event", "访视", "访视名称", "访视周期"],
    "visit_date": ["visit date", "assessment date", "collection date", "date", "访视日期", "检查日期", "采样日期"],
    "screening_date": ["screening date", "筛选日期", "icf签署日期", "informed consent date"],
    "randomization_date": ["randomization date", "随机日期", "randomized date"],
    "first_dose_date": ["first dose date", "首次给药日期", "first dosing date", "首次用药日期"],
    "withdrawal_date": ["withdrawal date", "退出日期", "end of study date", "脱落日期", "终止日期"],
    "subject_status": ["subject status", "status", "受试者状态", "受试者结局", "disposition"],
    "ae_term": ["ae term", "adverse event", "不良事件名称", "ae名称"],
    "ae_start_date": ["ae start date", "不良事件开始日期", "发生日期"],
    "ae_end_date": ["ae end date", "不良事件结束日期", "结束日期"],
    "ae_severity": ["severity", "intensity", "ctcae grade", "严重程度", "分级"],
    "is_sae": ["sae", "serious", "是否严重不良事件", "serious event"],
    "drug_name": ["drug name", "药物名称", "study drug", "imp name"],
    "dose_value": ["dose", "剂量", "给药剂量"],
    "dose_date": ["dosing date", "给药日期", "administration date"],
}

MODULE_HINTS = {
    "Laboratory": ["lab", "实验室", "化验", "检验", "alt", "ast", "hba1c", "fpg"],
    "Adverse Event": ["ae", "adverse", "不良事件", "sae", "aesi"],
    "Concomitant Medication": ["cm", "conmed", "合并用药"],
    "Drug Administration": ["dose", "drug", "给药", "exposure", "administration", "ip"],
    "Visit": ["visit", "访视", "folder", "event"],
    "Subject/Enrollment": ["subject", "demography", "screen", "random", "受试者", "人口学", "筛选", "随机"],
    "Disposition": ["status", "withdraw", "退出", "终止", "disposition"],
}


def norm_text(value: Any) -> str:
    return re.sub(r"\s+", " ", str(value).strip().lower())


def infer_module(sheet_name: str, columns: List[str]) -> str:
    text = f"{sheet_name} {' '.join(map(str, columns[:50]))}".lower()
    scores = {module: sum(1 for hint in hints if hint in text) for module, hints in MODULE_HINTS.items()}
    best = max(scores, key=scores.get)
    return best if scores[best] > 0 else "Unknown"


def score_mapping(col: str, synonyms: List[str]) -> int:
    c = norm_text(col)
    best = 0
    for syn in synonyms:
        s = norm_text(syn)
        if c == s:
            best = max(best, 100)
        elif s in c or c in s:
            best = max(best, 85)
        else:
            c_tokens = set(re.split(r"[^a-zA-Z0-9\u4e00-\u9fff]+", c))
            s_tokens = set(re.split(r"[^a-zA-Z0-9\u4e00-\u9fff]+", s))
            overlap = len(c_tokens & s_tokens)
            if overlap:
                best = max(best, 50 + overlap * 10)
    return best


def recommend_mapping(sheet_dict: Dict[str, pd.DataFrame]) -> pd.DataFrame:
    rows = []
    for sheet_name, df in sheet_dict.items():
        for std_field, synonyms in FIELD_SYNONYMS.items():
            best_col, best_score = None, 0
            for col in df.columns:
                score = score_mapping(str(col), synonyms)
                if score > best_score:
                    best_col, best_score = str(col), score
            if best_col and best_score >= 50:
                rows.append({
                    "standard_field_name": std_field,
                    "sheet_name": sheet_name,
                    "original_column_name": best_col,
                    "mapping_confidence": "高" if best_score >= 90 else "中" if best_score >= 70 else "低",
                    "score": best_score,
                    "needs_review": best_score < 90,
                })
    if not rows:
        return pd.DataFrame(columns=["standard_field_name", "sheet_name", "original_column_name", "mapping_confidence", "score", "needs_review"])
    return pd.DataFrame(rows).sort_values(["standard_field_name", "sheet_name", "score"], ascending=[True, True, False]).drop_duplicates(["standard_field_name", "sheet_name"], keep="first").reset_index(drop=True)


def get_local_col(mapping_df: pd.DataFrame, sheet_name: str, std_field: str):
    if mapping_df.empty:
        return None
    matched = mapping_df[(mapping_df["sheet_name"] == sheet_name) & (mapping_df["standard_field_name"] == std_field)]
    if matched.empty:
        return None
    return str(matched.sort_values("score", ascending=False).iloc[0]["original_column_name"])


def to_dt(value):
    if pd.isna(value) or value == "":
        return pd.NaT
    try:
        return pd.to_datetime(value)
    except Exception:
        return pd.NaT


def read_protocol_text(uploaded_file) -> str:
    if uploaded_file is None:
        return ""
    name = uploaded_file.name.lower()
    data = uploaded_file.read()
    uploaded_file.seek(0)
    if name.endswith(".pdf"):
        if fitz is None:
            return "未安装PyMuPDF，无法解析PDF。"
        doc = fitz.open(stream=data, filetype="pdf")
        return "\n".join(page.get_text() for page in doc)
    if name.endswith(".docx"):
        from docx import Document as DocxDocument
        doc = DocxDocument(io.BytesIO(data))
        return "\n".join(p.text for p in doc.paragraphs)
    return data.decode("utf-8", errors="ignore")


def run_basic_audit(sheet_dict: Dict[str, pd.DataFrame], mapping_df: pd.DataFrame) -> pd.DataFrame:
    findings = []

    def add(sheet, row_no, subject, category, title, field, value, risk, evidence):
        findings.append({
            "Finding ID": f"F{len(findings)+1:04d}",
            "Sheet": sheet,
            "Row": row_no,
            "Subject ID": subject,
            "Category": category,
            "Title": title,
            "Field": field,
            "Current Value": value,
            "Risk Level": risk,
            "Evidence": evidence,
        })

    required = ["subject_id", "visit_name", "visit_date"]
    mapped_fields = set(mapping_df["standard_field_name"].tolist()) if not mapping_df.empty else set()
    for field in required:
        if field not in mapped_fields:
            add("", "", "", "Data Integrity", f"关键字段未识别：{field}", field, "", "Major", "字段映射表中未找到对应字段。")

    for sheet_name, df in sheet_dict.items():
        if df.empty:
            continue
        subject_col = get_local_col(mapping_df, sheet_name, "subject_id")
        visit_col = get_local_col(mapping_df, sheet_name, "visit_name")
        visit_date_col = get_local_col(mapping_df, sheet_name, "visit_date")
        screening_col = get_local_col(mapping_df, sheet_name, "screening_date")
        random_col = get_local_col(mapping_df, sheet_name, "randomization_date")
        first_dose_col = get_local_col(mapping_df, sheet_name, "first_dose_date")
        withdrawal_col = get_local_col(mapping_df, sheet_name, "withdrawal_date")
        status_col = get_local_col(mapping_df, sheet_name, "subject_status")
        ae_severity_col = get_local_col(mapping_df, sheet_name, "ae_severity")
        is_sae_col = get_local_col(mapping_df, sheet_name, "is_sae")

        if subject_col and subject_col in df.columns:
            missing_subject = df[df[subject_col].isna() | (df[subject_col].astype(str).str.strip() == "")]
            for idx, row in missing_subject.head(200).iterrows():
                add(sheet_name, idx + 2, "", "Data Integrity", "受试者编号缺失", subject_col, "", "Major", f"{sheet_name}.{subject_col}为空。")

        if visit_col and visit_date_col and visit_col in df.columns and visit_date_col in df.columns:
            missing_visit_date = df[df[visit_col].notna() & (df[visit_col].astype(str).str.strip() != "") & (df[visit_date_col].isna() | (df[visit_date_col].astype(str).str.strip() == ""))]
            for idx, row in missing_visit_date.head(200).iterrows():
                subject = row.get(subject_col, "") if subject_col else ""
                add(sheet_name, idx + 2, subject, "Visit", "访视日期缺失", visit_date_col, "", "Major", "存在访视记录但访视日期为空。")

        if subject_col and visit_col and subject_col in df.columns and visit_col in df.columns:
            dup = df.groupby([subject_col, visit_col]).size().reset_index(name="n")
            dup = dup[dup["n"] > 1]
            for _, row in dup.head(200).iterrows():
                add(sheet_name, "", row[subject_col], "Visit", "同一受试者同一访视重复记录", visit_col, f"count={row['n']}", "Minor", "需确认是否为重复录入、重测或非计划访视。")

        if subject_col and screening_col and random_col and screening_col in df.columns and random_col in df.columns:
            temp = df[[subject_col, screening_col, random_col]].copy()
            temp["_screen"] = temp[screening_col].map(to_dt)
            temp["_random"] = temp[random_col].map(to_dt)
            bad = temp[temp["_screen"].notna() & temp["_random"].notna() & (temp["_random"] < temp["_screen"])]
            for idx, row in bad.head(200).iterrows():
                add(sheet_name, idx + 2, row.get(subject_col, ""), "Eligibility", "随机日期早于筛选日期", f"{screening_col}/{random_col}", f"{row[screening_col]} / {row[random_col]}", "Major", "筛选-随机主线时间顺序冲突。")

        if subject_col and first_dose_col and random_col and first_dose_col in df.columns and random_col in df.columns:
            temp = df[[subject_col, first_dose_col, random_col]].copy()
            temp["_dose"] = temp[first_dose_col].map(to_dt)
            temp["_random"] = temp[random_col].map(to_dt)
            bad = temp[temp["_dose"].notna() & temp["_random"].notna() & (temp["_dose"] < temp["_random"])]
            for idx, row in bad.head(200).iterrows():
                add(sheet_name, idx + 2, row.get(subject_col, ""), "Dosing", "首次给药日期早于随机日期", f"{first_dose_col}/{random_col}", f"{row[first_dose_col]} / {row[random_col]}", "Major", "随机-给药主线时间顺序冲突。")

        if subject_col and status_col and withdrawal_col and status_col in df.columns and withdrawal_col in df.columns:
            status = df[status_col].astype(str).str.lower()
            bad = df[status.apply(lambda x: any(k in x for k in ["withdraw", "discontinue", "退出", "终止", "脱落"])) & (df[withdrawal_col].isna() | (df[withdrawal_col].astype(str).str.strip() == ""))]
            for idx, row in bad.head(200).iterrows():
                add(sheet_name, idx + 2, row.get(subject_col, ""), "Disposition", "退出/终止状态但退出日期缺失", withdrawal_col, "", "Major", "受试者状态与退出日期不一致。")

        if subject_col and ae_severity_col and is_sae_col and ae_severity_col in df.columns and is_sae_col in df.columns:
            sev = df[ae_severity_col].astype(str).str.lower()
            sae = df[is_sae_col].astype(str).str.lower()
            bad = df[sev.apply(lambda x: any(k in x for k in ["grade 4", "grade 5", "4", "5", "serious", "severe", "重度"])) & ~sae.apply(lambda x: x in ["yes", "y", "true", "1", "是"])]
            for idx, row in bad.head(200).iterrows():
                add(sheet_name, idx + 2, row.get(subject_col, ""), "Safety", "AE严重程度与SAE标记需复核", f"{ae_severity_col}/{is_sae_col}", f"{row[ae_severity_col]} / {row[is_sae_col]}", "Major", "AE严重程度较高但SAE字段未标记为是，需医学与数据复核。")

    return pd.DataFrame(findings)


def build_word_report(findings_df: pd.DataFrame, sheet_meta_df: pd.DataFrame, mapping_df: pd.DataFrame, protocol_summary: str) -> bytes:
    doc = Document()
    doc.add_heading("EDC智能稽查与RBQM风险识别报告", 0)
    doc.add_paragraph(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_heading("一、项目资料读取概况", level=1)
    doc.add_paragraph(f"读取Sheet数量：{len(sheet_meta_df)}")
    doc.add_paragraph(f"识别字段映射数量：{len(mapping_df)}")
    if protocol_summary:
        doc.add_heading("二、方案文本摘要", level=1)
        doc.add_paragraph(protocol_summary[:1500])
    doc.add_heading("三、主要发现", level=1)
    if findings_df.empty:
        doc.add_paragraph("本次MVP规则未检出明确数据问题。仍建议结合方案规则、SDV/SDR策略和医学判断进一步复核。")
    else:
        table = doc.add_table(rows=1, cols=6)
        hdr = table.rows[0].cells
        for i, title in enumerate(["Finding ID", "Sheet", "Subject ID", "Title", "Risk Level", "Evidence"]):
            hdr[i].text = title
        for _, r in findings_df.head(200).iterrows():
            cells = table.add_row().cells
            for i, title in enumerate(["Finding ID", "Sheet", "Subject ID", "Title", "Risk Level", "Evidence"]):
                cells[i].text = str(r.get(title, ""))
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


with st.sidebar:
    st.header("上传文件")
    edc_file = st.file_uploader("上传EDC导出Excel", type=["xlsx", "xls"])
    protocol_file = st.file_uploader("上传方案/规则文件（可选）", type=["pdf", "docx", "txt", "md"])

if not edc_file:
    st.info("请先上传EDC Excel文件。系统会自动读取全部Sheet并生成字段映射和基础稽查发现。")
    st.stop()

try:
    sheet_dict = pd.read_excel(edc_file, sheet_name=None)
except Exception as exc:
    st.error(f"Excel读取失败：{exc}")
    st.stop()

sheet_meta_rows = []
for sheet_name, df in sheet_dict.items():
    sheet_meta_rows.append({
        "sheet_name": sheet_name,
        "row_count": len(df),
        "column_count": len(df.columns),
        "has_data": not df.empty,
        "inferred_module": infer_module(sheet_name, [str(c) for c in df.columns]),
    })
sheet_meta_df = pd.DataFrame(sheet_meta_rows)

mapping_df = recommend_mapping(sheet_dict)
protocol_text = read_protocol_text(protocol_file) if protocol_file else ""
findings_df = run_basic_audit(sheet_dict, mapping_df)

metric1, metric2, metric3, metric4 = st.columns(4)
metric1.metric("Sheet数量", len(sheet_meta_df))
metric2.metric("总行数", int(sheet_meta_df["row_count"].sum()))
metric3.metric("字段映射", len(mapping_df))
metric4.metric("发现数量", len(findings_df))

tab1, tab2, tab3, tab4, tab5 = st.tabs(["Sheet概况", "数据预览", "字段映射", "稽查发现", "报告导出"])

with tab1:
    st.subheader("Sheet读取概况")
    st.dataframe(sheet_meta_df, use_container_width=True)

with tab2:
    selected_sheet = st.selectbox("选择Sheet", list(sheet_dict.keys()))
    st.dataframe(sheet_dict[selected_sheet].head(100), use_container_width=True)

with tab3:
    st.subheader("自动推荐字段映射")
    if mapping_df.empty:
        st.warning("未识别到稳定字段映射，请检查表头命名或后续补充字段词典。")
    else:
        edited_mapping = st.data_editor(mapping_df, use_container_width=True, num_rows="dynamic")
        mapping_df = edited_mapping

with tab4:
    st.subheader("基础稽查发现")
    if findings_df.empty:
        st.success("当前MVP规则未发现明确问题。")
    else:
        st.dataframe(findings_df, use_container_width=True)
        st.download_button("下载Findings CSV", findings_df.to_csv(index=False).encode("utf-8-sig"), file_name="edc_findings.csv", mime="text/csv")

with tab5:
    st.subheader("生成Word稽查报告")
    protocol_summary = protocol_text[:1500] if protocol_text else ""
    report_bytes = build_word_report(findings_df, sheet_meta_df, mapping_df, protocol_summary)
    st.download_button("下载Word报告", report_bytes, file_name="EDC智能稽查报告.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
