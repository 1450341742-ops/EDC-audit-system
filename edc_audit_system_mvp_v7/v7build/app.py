import io
import re
from dataclasses import dataclass
from datetime import datetime
from typing import Any, Dict, List, Optional, Tuple

import pandas as pd
import streamlit as st
from docx import Document
from docx.shared import Pt

try:
    import fitz  # PyMuPDF
except Exception:
    fitz = None


st.set_page_config(page_title="EDC智能稽查与RBQM风险识别系统", layout="wide")


FIELD_SYNONYMS = {
    "subject_id": ["subject id", "subject no", "screening no", "randomization no", "受试者编号", "受试者id", "受试者", "随机号", "筛选号"],
    "site_id": ["site", "site no", "center", "center no", "中心号", "中心编号", "研究中心"],
    "visit_name": ["visit", "visit name", "visit label", "folder", "event", "instance", "访视", "访视名称", "访视周期"],
    "visit_date": ["visit date", "assessment date", "collection date", "date", "访视日期", "检查日期", "采样日期", "随访日期"],
    "screening_date": ["screening date", "筛选日期", "筛选访问日期", "icf签署日期", "informed consent date"],
    "randomization_date": ["randomization date", "随机日期", "randomized date"],
    "first_dose_date": ["first dose date", "首次给药日期", "first dosing date", "首次用药日期"],
    "last_dose_date": ["last dose date", "末次给药日期", "last dosing date", "末次用药日期"],
    "withdrawal_date": ["withdrawal date", "退出日期", "end of study date", "脱落日期", "终止日期"],
    "subject_status": ["subject status", "status", "受试者状态", "受试者结局", "disposition"],
    "lab_test_name": ["test name", "lab test", "测定项目", "检验项目", "实验室项目"],
    "lab_result": ["result", "检查结果", "测定值", "结果值", "value"],
    "lab_unit": ["unit", "单位"],
    "lab_uln": ["uln", "upper limit", "范围上限", "正常上限"],
    "lab_lln": ["lln", "lower limit", "范围下限", "正常下限"],
    "ae_term": ["ae term", "adverse event", "不良事件名称", "ae名称"],
    "ae_start_date": ["ae start date", "不良事件开始日期", "发生日期"],
    "ae_end_date": ["ae end date", "不良事件结束日期", "结束日期"],
    "ae_severity": ["severity", "intensity", "ctcae grade", "严重程度", "分级"],
    "is_sae": ["sae", "serious", "是否严重不良事件", "serious event"],
    "is_aesi": ["aesi", "special interest", "特别关注不良事件"],
    "drug_name": ["drug name", "药物名称", "study drug", "imp name"],
    "dose_value": ["dose", "剂量", "给药剂量"],
    "dose_date": ["dosing date", "给药日期", "administration date"],
    "cm_drug_name": ["conmed", "cm drug", "合并用药", "药物名称"],
    "cm_start_date": ["cm start date", "合并用药开始日期", "开始日期"],
    "cm_end_date": ["cm end date", "合并用药结束日期", "结束日期"],
}

MODULE_HINTS = {
    "Laboratory": ["lab", "实验室", "化验", "检验", "hba1c", "alt", "ast", "fpg"],
    "Adverse Event": ["ae", "adverse", "不良事件", "sae", "aesi"],
    "Concomitant Medication": ["cm", "conmed", "合并用药"],
    "Drug Administration": ["dose", "drug", "给药", "exposure", "administration", "ip"],
    "Visit": ["visit", "访视", "folder", "event"],
    "Subject/Enrollment": ["subject", "demography", "screen", "random", "受试者", "人口学", "筛选", "随机"],
    "Disposition": ["status", "withdraw", "退出", "终止", "disposition"],
}

SCREENING_KEYWORDS = [
    "screen", "screening", "筛选", "筛查", "scr", "visit 1", "v1", "sv"
]


@dataclass
class SheetMeta:
    sheet_name: str
    row_count: int
    column_count: int
    has_data: bool
    inferred_module: str
    read_status: str


def norm_text(x: Any) -> str:
    return re.sub(r"\s+", " ", str(x).strip().lower())


def infer_module(sheet_name: str, columns: List[str]) -> str:
    text = f"{sheet_name} {' '.join(map(str, columns[:40]))}".lower()
    scores = {}
    for mod, hints in MODULE_HINTS.items():
        scores[mod] = sum(1 for h in hints if h.lower() in text)
    best = max(scores, key=scores.get)
    return best if scores[best] > 0 else "Unknown"


def score_mapping(col: str, targets: List[str]) -> int:
    c = norm_text(col)
    score = 0
    for t in targets:
        t2 = norm_text(t)
        if c == t2:
            score = max(score, 100)
        elif t2 in c or c in t2:
            score = max(score, 85)
        else:
            toks_c = set(re.split(r"[^a-zA-Z0-9\u4e00-\u9fff]+", c))
            toks_t = set(re.split(r"[^a-zA-Z0-9\u4e00-\u9fff]+", t2))
            overlap = len(toks_c & toks_t)
            if overlap:
                score = max(score, 50 + overlap * 10)
    return score


def recommend_field_mapping(sheet_dict: Dict[str, pd.DataFrame]) -> pd.DataFrame:
    rows = []
    for sheet_name, df in sheet_dict.items():
        cols = [str(c) for c in df.columns]
        for std_field, syns in FIELD_SYNONYMS.items():
            best_col, best_score = None, 0
            for col in cols:
                sc = score_mapping(col, syns)
                if sc > best_score:
                    best_col, best_score = col, sc
            if best_col and best_score >= 50:
                rows.append({
                    "standard_field_name": std_field,
                    "sheet_name": sheet_name,
                    "original_column_name": best_col,
                    "mapping_confidence": "高" if best_score >= 90 else "中" if best_score >= 70 else "低",
                    "score": best_score,
                    "is_decode": "decode" in norm_text(best_col) or "译码" in best_col,
                    "needs_review": best_score < 90,
                })
    if not rows:
        return pd.DataFrame(columns=["standard_field_name", "sheet_name", "original_column_name", "mapping_confidence", "score", "is_decode", "needs_review"])
    out = pd.DataFrame(rows).sort_values(["standard_field_name", "sheet_name", "score"], ascending=[True, True, False])
    # 保留“每个标准字段在每个Sheet中的最佳列”，而不是全局只保留第一个Sheet
    out = out.drop_duplicates(["standard_field_name", "sheet_name"], keep="first")
    return out.reset_index(drop=True)


def read_protocol_text(file) -> str:
    name = file.name.lower()
    data = file.read()
    file.seek(0)
    if name.endswith(".pdf"):
        if not fitz:
            return "[未安装 PyMuPDF，无法解析 PDF 文本]"
        doc = fitz.open(stream=data, filetype="pdf")
        return "\n".join(page.get_text() for page in doc)
    if name.endswith(".docx"):
        from docx import Document as DocxDocument
        doc = DocxDocument(io.BytesIO(data))
        return "\n".join(p.text for p in doc.paragraphs)
    return data.decode("utf-8", errors="ignore")


def extract_protocol_meta(text: str) -> Dict[str, Any]:
    version = re.search(r"v(?:ersion)?\s*([0-9]+(?:\.[0-9]+)?)", text, re.I)
    date = re.search(r"(20\d{2}[-/年]\d{1,2}[-/月]\d{1,2})", text)
    protocol = re.search(r"([A-Z]{2,}[A-Z0-9\-_/]+)", text)
    return {
        "protocol_name": protocol.group(1) if protocol else "未识别",
        "version": version.group(1) if version else "未识别",
        "version_date": date.group(1) if date else "未识别",
    }


def heuristic_rule_extract(text: str) -> pd.DataFrame:
    rules = []
    joined = "\n".join(text.split("\n")[:3000])
    patterns = [
        ("Eligibility", r"BMI[^\n]{0,30}?([<>]=?)\s*([0-9]+(?:\.[0-9]+)?)", "BMI threshold"),
        ("Laboratory", r"HbA1c[^\n]{0,40}?([<>]=?)\s*([0-9]+(?:\.[0-9]+)?)", "HbA1c threshold"),
        ("Laboratory", r"FPG[^\n]{0,40}?([<>]=?)\s*([0-9]+(?:\.[0-9]+)?)", "FPG threshold"),
        ("Laboratory", r"ALT[^\n]{0,40}?([<>]=?)\s*([0-9]+(?:\.[0-9]+)?)", "ALT threshold"),
        ("Laboratory", r"AST[^\n]{0,40}?([<>]=?)\s*([0-9]+(?:\.[0-9]+)?)", "AST threshold"),
        ("Safety", r"SAE|serious adverse event|严重不良事件", "SAE requirement"),
        ("Safety", r"AESI|special interest|特别关注不良事件", "AESI requirement"),
        ("Visit", r"window|窗口", "Visit window"),
        ("Dosing", r"dose|给药|剂量", "Dosing rule"),
    ]
    rule_id = 1
    for category, pat, name in patterns:
        for m in re.finditer(pat, joined, re.I):
            op = m.group(1) if m.lastindex and m.lastindex >= 1 else ""
            value = m.group(2) if m.lastindex and m.lastindex >= 2 else ""
            snippet = joined[max(0, m.start()-80): m.end()+80].replace("\n", " ")
            rules.append({
                "rule_id": f"R{rule_id:03d}",
                "category": category,
                "rule_name": name,
                "rule_text": snippet,
                "required_field": name.split(" threshold")[0] if "threshold" in name else "",
                "operator": op,
                "threshold": value,
                "risk_level": "Major" if category in ["Eligibility", "Safety"] else "Minor",
                "protocol_reference": snippet[:120],
                "medical_review_required": category == "Safety",
            })
            rule_id += 1
    if not rules:
        rules.append({
            "rule_id": "R001", "category": "General", "rule_name": "Manual rule review needed", "rule_text": "未从文本中稳定抽取到结构化规则，请人工编辑补充。",
            "required_field": "", "operator": "", "threshold": "", "risk_level": "Observation", "protocol_reference": "", "medical_review_required": True,
        })
    return pd.DataFrame(rules)


def to_dt(x):
    if pd.isna(x) or x == "":
        return pd.NaT
    try:
        return pd.to_datetime(x)
    except Exception:
        return pd.NaT


def get_mapped_info(mapping_df: pd.DataFrame, std_field: str) -> Optional[Tuple[str, str]]:
    m = mapping_df[mapping_df["standard_field_name"] == std_field]
    if m.empty:
        return None
    row = m.sort_values(["score", "mapping_confidence"], ascending=[False, True]).iloc[0]
    return row["sheet_name"], row["original_column_name"]


def get_mapped_infos(mapping_df: pd.DataFrame, std_field: str) -> List[Tuple[str, str]]:
    m = mapping_df[mapping_df["standard_field_name"] == std_field]
    if m.empty:
        return []
    m = m.sort_values(["sheet_name", "score"], ascending=[True, False]).drop_duplicates(["sheet_name"], keep="first")
    return [(str(r["sheet_name"]), str(r["original_column_name"])) for _, r in m.iterrows()]


def get_mapped_series(sheet_dict: Dict[str, pd.DataFrame], mapping_df: pd.DataFrame, std_field: str) -> Optional[Tuple[str, pd.Series]]:
    info = get_mapped_info(mapping_df, std_field)
    if not info:
        return None
    sheet, col = info
    if sheet in sheet_dict and col in sheet_dict[sheet].columns:
        return sheet, sheet_dict[sheet][col]
    return None


def get_local_col(mapping_df: pd.DataFrame, sheet_name: str, std_field: str) -> Optional[str]:
    m = mapping_df[(mapping_df["standard_field_name"] == std_field) & (mapping_df["sheet_name"] == sheet_name)]
    if m.empty:
        return None
    row = m.sort_values(["score"], ascending=[False]).iloc[0]
    return str(row["original_column_name"])


def screening_mask(df: pd.DataFrame, sheet_name: str, mapping_df: pd.DataFrame) -> pd.Series:
    mask = pd.Series([True] * len(df), index=df.index)

    # Priority 1: mapped visit_name on same sheet
    local_visit_col = get_local_col(mapping_df, sheet_name, "visit_name")
    if local_visit_col and local_visit_col in df.columns:
        vals = df[local_visit_col].astype(str).str.lower()
        scr = vals.apply(lambda x: any(k in x for k in SCREENING_KEYWORDS))
        if scr.any():
            return scr.fillna(False)

    # Priority 2: infer by local columns containing visit/folder/event
    candidate_cols = [
        c for c in df.columns
        if any(k in norm_text(c) for k in ["visit", "folder", "event", "instance", "访视", "周期", "阶段"])
    ]
    for c in candidate_cols:
        vals = df[c].astype(str).str.lower()
        scr = vals.apply(lambda x: any(k in x for k in SCREENING_KEYWORDS))
        if scr.any():
            return scr.fillna(False)

    # Priority 3: if sheet name itself indicates screening, keep all rows
    if any(k in norm_text(sheet_name) for k in SCREENING_KEYWORDS):
        return mask

    # Priority 4: if no visit marker found, keep all rows to avoid only first sheet findings
    return mask


def build_screening_sheet_dict(sheet_dict: Dict[str, pd.DataFrame], mapping_df: pd.DataFrame) -> Dict[str, pd.DataFrame]:
    out = {}
    for s, df in sheet_dict.items():
        if df.empty:
            out[s] = df.copy()
            continue
        m = screening_mask(df, s, mapping_df)
        sub = df[m].copy()
        out[s] = sub
    return out


def restricted_mode(sheet_meta_df: pd.DataFrame, mapping_df: pd.DataFrame) -> Tuple[bool, List[str]]:
    reasons = []
    if len(sheet_meta_df) <= 1:
        reasons.append("当前仅识别到 1 个 Sheet，属于高风险读取受限场景")
    key_fields = {"subject_id", "visit_name", "visit_date", "randomization_date", "first_dose_date", "withdrawal_date", "ae_term", "dose_date"}
    got = set(mapping_df["standard_field_name"].tolist()) if not mapping_df.empty else set()
    missing = key_fields - got
    if len(missing) >= 4:
        reasons.append(f"关键主线字段缺失较多：{', '.join(sorted(missing))}")
    return (len(reasons) > 0, reasons)




def build_sheet_coverage(sheet_meta_df: pd.DataFrame, filtered_sheet_meta_df: Optional[pd.DataFrame], mapping_df: pd.DataFrame, findings_df: Optional[pd.DataFrame] = None) -> pd.DataFrame:
    rows = []
    findings_df = findings_df if findings_df is not None else pd.DataFrame()
    filt_lookup = {}
    if filtered_sheet_meta_df is not None and not filtered_sheet_meta_df.empty and 'sheet_name' in filtered_sheet_meta_df.columns:
        for _, r in filtered_sheet_meta_df.iterrows():
            filt_lookup[str(r['sheet_name'])] = int(r.get('row_count', 0) or 0)
    for _, r in sheet_meta_df.iterrows():
        s = str(r['sheet_name'])
        map_count = int((mapping_df['sheet_name'].astype(str) == s).sum()) if not mapping_df.empty and 'sheet_name' in mapping_df.columns else 0
        finding_count = int((findings_df['Sheet'].astype(str) == s).sum()) if (findings_df is not None and not findings_df.empty and 'Sheet' in findings_df.columns) else 0
        rows.append({
            'Sheet': s,
            'Rows Read': int(r.get('row_count', 0) or 0),
            'Columns': int(r.get('column_count', 0) or 0),
            'Module': r.get('inferred_module', ''),
            'Mapped Fields': map_count,
            'Rows After Screening': filt_lookup.get(s, int(r.get('row_count', 0) or 0)),
            'Findings Count': finding_count,
            'Coverage Status': 'Has findings' if finding_count > 0 else ('Mapped but no rule hit' if map_count > 0 else 'Read only / unmapped')
        })
    return pd.DataFrame(rows)

def run_audit(sheet_dict: Dict[str, pd.DataFrame], mapping_df: pd.DataFrame, rules_df: pd.DataFrame) -> Tuple[pd.DataFrame, pd.DataFrame, pd.DataFrame, pd.DataFrame]:
    findings = []

    def add_finding(subject_id: Any, site: Any, visit: Any, category: str, title: str, description: str,
                    sheet: str, field: str, current_value: Any, protocol_requirement: str,
                    judgement: str, risk_level: str, evidence: str, need_medical_review: str = "No"):
        findings.append({
            "Finding ID": f"F{len(findings)+1:04d}",
            "Subject ID": "" if pd.isna(subject_id) else subject_id,
            "Site": "" if pd.isna(site) else site,
            "Visit": "" if pd.isna(visit) else visit,
            "Category": category,
            "Title": title,
            "Description": description,
            "Sheet": sheet,
            "Field": field,
            "Current Value": current_value,
            "Protocol Requirement": protocol_requirement,
            "Judgement": judgement,
            "Risk Level": risk_level,
            "Evidence": evidence,
            "Need Medical Review": need_medical_review,
        })

    # 全局关键字段缺失提示
    for f in ["subject_id", "visit_name", "visit_date"]:
        if mapping_df[mapping_df["standard_field_name"] == f].empty:
            add_finding("", "", "", "Data Integrity", f"关键字段缺失：{f}",
                        f"当前未完成标准字段 {f} 的稳定映射，无法支撑完整核查。",
                        "", "", "", "关键主线字段需可识别",
                        "Unable to assess / Review needed", "Major", "字段映射表中未找到对应映射", "No")

    # 逐Sheet多模块核查
    for sheet_name, df in sheet_dict.items():
        if df is None or df.empty:
            continue
        subject_col = get_local_col(mapping_df, sheet_name, "subject_id")
        site_col = get_local_col(mapping_df, sheet_name, "site_id")
        visit_col = get_local_col(mapping_df, sheet_name, "visit_name")
        visit_date_col = get_local_col(mapping_df, sheet_name, "visit_date")
        screening_date_col = get_local_col(mapping_df, sheet_name, "screening_date")
        randomization_date_col = get_local_col(mapping_df, sheet_name, "randomization_date")
        withdrawal_date_col = get_local_col(mapping_df, sheet_name, "withdrawal_date")
        subject_status_col = get_local_col(mapping_df, sheet_name, "subject_status")

        # Visit / missing date / duplicate visit
        if subject_col and visit_col and subject_col in df.columns and visit_col in df.columns:
            temp = pd.DataFrame({
                "subject": df[subject_col].astype(str),
                "site": df[site_col] if site_col and site_col in df.columns else "",
                "visit": df[visit_col].astype(str),
                "visit_date": df[visit_date_col] if visit_date_col and visit_date_col in df.columns else pd.NaT,
            })
            dup = temp.groupby(["subject", "visit"]).size().reset_index(name="n")
            dup = dup[dup["n"] > 1]
            for _, r in dup.iterrows():
                add_finding(r["subject"], "", r["visit"], "Visit",
                            "Duplicate visit record",
                            "同一受试者在同一Sheet存在重复访视记录，需核对是否为重复录入或非计划访视。",
                            sheet_name, visit_col, f"count={r['n']}",
                            "Each visit should be uniquely attributable or clearly explained",
                            "Review needed", "Minor", f"{sheet_name}: duplicate subject-visit pairs", "No")
            if visit_date_col and visit_date_col in df.columns:
                miss = temp[temp["visit_date"].isna() | (temp["visit_date"].astype(str).str.strip() == "")]
                for _, r in miss.head(200).iterrows():
                    add_finding(r["subject"], r["site"], r["visit"], "Visit",
                                "Missing Visit Date",
                                "访视记录存在但访视日期缺失，影响窗口判断与主线逻辑。",
                                sheet_name, visit_date_col, "", "Visit Date should be recorded",
                                "Not compliant", "Major", f"{sheet_name}.{visit_date_col} is blank", "No")

        # Eligibility / status based local date checks
        if subject_col and subject_col in df.columns and screening_date_col and randomization_date_col and screening_date_col in df.columns and randomization_date_col in df.columns:
            temp = pd.DataFrame({
                "subject": df[subject_col].astype(str),
                "site": df[site_col] if site_col and site_col in df.columns else "",
                "screening": df[screening_date_col].map(to_dt),
                "randomization": df[randomization_date_col].map(to_dt),
            })
            bad = temp[(temp["screening"].notna()) & (temp["randomization"].notna()) & (temp["randomization"] < temp["screening"])]
            for _, r in bad.iterrows():
                add_finding(r["subject"], r["site"], "Screening", "Eligibility",
                            "Randomization earlier than Screening",
                            "随机日期早于筛选日期，存在入组主线逻辑冲突。",
                            sheet_name, f"{screening_date_col} / {randomization_date_col}",
                            f"{r['screening']} / {r['randomization']}",
                            "Screening should occur before Randomization",
                            "Not compliant", "Major", "Local sheet date order conflict", "No")

        if subject_col and subject_status_col and subject_col in df.columns and subject_status_col in df.columns:
            if withdrawal_date_col and withdrawal_date_col in df.columns:
                temp = pd.DataFrame({
                    "subject": df[subject_col].astype(str),
                    "site": df[site_col] if site_col and site_col in df.columns else "",
                    "status": df[subject_status_col].astype(str),
                    "wd": df[withdrawal_date_col],
                })
                withdrawn_words = ["withdraw", "退出", "脱落", "终止", "discontinue"]
                bad = temp[temp["status"].str.lower().apply(lambda x: any(w in x for w in withdrawn_words)) & (temp["wd"].isna() | (temp["wd"].astype(str).str.strip() == ""))]
                for _, r in bad.iterrows():
                    add_finding(r["subject"], r["site"], "Disposition", "Data Integrity",
                                "Withdrawal status without Withdrawal Date",
                                "受试者状态提示已退出/脱落，但退出日期缺失。",
                                sheet_name, f"{subject_status_col} / {withdrawal_date_col}",
                                f"{r['status']} / {r['wd']}",
                                "Withdrawal Date should be recorded for withdrawn subject",
                                "Review needed", "Major", "Disposition status conflicts with missing withdrawal date", "No")

        # Laboratory rules across all mapped lab sheets
        lab_name_col = get_local_col(mapping_df, sheet_name, "lab_test_name")
        lab_result_col = get_local_col(mapping_df, sheet_name, "lab_result")
        lab_unit_col = get_local_col(mapping_df, sheet_name, "lab_unit")
        if subject_col and lab_name_col and lab_result_col and subject_col in df.columns and lab_name_col in df.columns and lab_result_col in df.columns:
            temp = pd.DataFrame({
                "subject": df[subject_col],
                "site": df[site_col] if site_col and site_col in df.columns else "",
                "visit": df[visit_col] if visit_col and visit_col in df.columns else "Screening",
                "lab": df[lab_name_col].astype(str),
                "result": pd.to_numeric(df[lab_result_col], errors="coerce"),
                "unit": df[lab_unit_col].astype(str) if lab_unit_col and lab_unit_col in df.columns else "",
            })
            # protocol threshold rules
            thr_rules = rules_df[(rules_df["threshold"].astype(str) != "") & (rules_df["required_field"].astype(str) != "")]
            for _, rule in thr_rules.iterrows():
                field = str(rule["required_field"]).strip().lower()
                try:
                    thr = float(rule["threshold"])
                except Exception:
                    continue
                op = str(rule["operator"])
                subset = temp[temp["lab"].str.lower().str.contains(field.lower(), na=False)]
                for _, r in subset.iterrows():
                    hit = False
                    if op == ">" and pd.notna(r["result"]) and r["result"] > thr:
                        hit = True
                    elif op == ">=" and pd.notna(r["result"]) and r["result"] >= thr:
                        hit = True
                    elif op == "<" and pd.notna(r["result"]) and r["result"] < thr:
                        hit = True
                    elif op == "<=" and pd.notna(r["result"]) and r["result"] <= thr:
                        hit = True
                    if hit:
                        add_finding(r["subject"], r["site"], r["visit"], "Laboratory",
                                    f"{rule['rule_name']} hit",
                                    f"实验室结果命中方案阈值规则：{field} {op} {thr}",
                                    sheet_name,
                                    f"{lab_name_col} / {lab_result_col}",
                                    f"{r['lab']}={r['result']}",
                                    rule["rule_text"],
                                    "Potential issue / Review needed",
                                    rule["risk_level"],
                                    str(rule["protocol_reference"]),
                                    "Yes" if bool(rule.get("medical_review_required", False)) else "No")
            # missing units for key labs
            if lab_unit_col and lab_unit_col in df.columns:
                key_mask = temp["lab"].str.lower().str.contains("hba1c|fpg|alt|ast|tg|lipase|amylase", na=False)
                miss = temp[key_mask & ((temp["unit"].isna()) | (temp["unit"].astype(str).str.strip() == ""))]
                for _, r in miss.head(100).iterrows():
                    add_finding(r["subject"], r["site"], r["visit"], "Laboratory",
                                "Missing laboratory unit",
                                "关键实验室项目存在结果值但单位缺失，可能影响阈值判断。",
                                sheet_name, f"{lab_name_col} / {lab_unit_col}", f"{r['lab']}={r['result']}",
                                "Key laboratory values should have unit", "Review needed", "Minor",
                                "Result present but unit missing", "No")

        # Dosing
        dose_date_col = get_local_col(mapping_df, sheet_name, "dose_date")
        dose_value_col = get_local_col(mapping_df, sheet_name, "dose_value")
        first_dose_local = get_local_col(mapping_df, sheet_name, "first_dose_date")
        if subject_col and dose_date_col and subject_col in df.columns and dose_date_col in df.columns:
            temp = pd.DataFrame({
                "subject": df[subject_col].astype(str),
                "site": df[site_col] if site_col and site_col in df.columns else "",
                "visit": df[visit_col] if visit_col and visit_col in df.columns else "Screening",
                "dose_date": df[dose_date_col].map(to_dt),
                "dose_value": df[dose_value_col] if dose_value_col and dose_value_col in df.columns else "",
                "first_dose": df[first_dose_local].map(to_dt) if first_dose_local and first_dose_local in df.columns else pd.NaT,
            })
            miss = temp[temp["dose_date"].isna()]
            for _, r in miss.head(100).iterrows():
                add_finding(r["subject"], r["site"], r["visit"], "Dosing",
                            "Missing Dosing Date", "存在给药记录但给药日期缺失。", sheet_name,
                            dose_date_col, r["dose_value"], "Dosing Date should be recorded",
                            "Not compliant", "Major", "Dose row without date", "No")
            if first_dose_local and first_dose_local in df.columns:
                bad = temp[(temp["dose_date"].notna()) & (temp["first_dose"].notna()) & (temp["dose_date"] < temp["first_dose"])]
                for _, r in bad.iterrows():
                    add_finding(r["subject"], r["site"], r["visit"], "Dosing",
                                "Dose earlier than First Dose Date",
                                "给药日期早于首次给药日期主线字段。", sheet_name,
                                f"{dose_date_col} / {first_dose_local}", f"{r['dose_date']} / {r['first_dose']}",
                                "Dose records should not precede First Dose Date",
                                "Review needed", "Major", "Local dosing timeline conflict", "No")

        # Safety across all mapped AE sheets
        ae_term_col = get_local_col(mapping_df, sheet_name, "ae_term")
        ae_sev_col = get_local_col(mapping_df, sheet_name, "ae_severity")
        is_sae_col = get_local_col(mapping_df, sheet_name, "is_sae")
        ae_end_col = get_local_col(mapping_df, sheet_name, "ae_end_date")
        if subject_col and ae_term_col and subject_col in df.columns and ae_term_col in df.columns:
            if ae_sev_col and is_sae_col and ae_sev_col in df.columns and is_sae_col in df.columns:
                temp = pd.DataFrame({
                    "subject": df[subject_col],
                    "site": df[site_col] if site_col and site_col in df.columns else "",
                    "visit": df[visit_col] if visit_col and visit_col in df.columns else "Screening",
                    "ae": df[ae_term_col].astype(str),
                    "sev": df[ae_sev_col].astype(str),
                    "sae": df[is_sae_col].astype(str),
                })
                severe_words = ["重", "grade 3", "grade 4", "severe", "iii", "iv"]
                no_words = ["否", "no", "n", "false", "0"]
                for _, r in temp.iterrows():
                    sev = norm_text(r["sev"])
                    sae = norm_text(r["sae"])
                    if any(w in sev for w in severe_words) and any(w in sae for w in no_words):
                        add_finding(r["subject"], r["site"], r["visit"], "Safety",
                                    "AE Severity vs SAE requires review",
                                    "检测到高等级/重度 AE 但 SAE 字段为否，需核对是否混淆严重程度与严重性。",
                                    sheet_name,
                                    f"{ae_sev_col} / {is_sae_col}",
                                    f"{r['sev']} / {r['sae']}",
                                    "AE Severity and SAE must be assessed separately",
                                    "Review needed",
                                    "Observation",
                                    "高等级 AE 与 SAE=No 并存，不等同于错误，但需重点复核",
                                    "Yes")
            if ae_end_col and ae_end_col in df.columns:
                temp2 = pd.DataFrame({
                    "subject": df[subject_col],
                    "site": df[site_col] if site_col and site_col in df.columns else "",
                    "visit": df[visit_col] if visit_col and visit_col in df.columns else "Screening",
                    "ae": df[ae_term_col].astype(str),
                    "ae_end": df[ae_end_col],
                })
                miss = temp2[temp2["ae"].astype(str).str.strip() != ""]
                miss = miss[miss["ae_end"].isna() | (miss["ae_end"].astype(str).str.strip() == "")]
                for _, r in miss.head(100).iterrows():
                    add_finding(r["subject"], r["site"], r["visit"], "Safety",
                                "AE End Date missing / follow-up pending",
                                "AE 已记录但结束日期缺失，需确认是否仍在随访中。",
                                sheet_name, ae_end_col, r["ae"], "AE should be followed to outcome where applicable",
                                "Review needed", "Minor", "AE term present but AE End Date blank", "Yes")

        # CM
        cm_name_col = get_local_col(mapping_df, sheet_name, "cm_drug_name")
        cm_start_col = get_local_col(mapping_df, sheet_name, "cm_start_date")
        cm_end_col = get_local_col(mapping_df, sheet_name, "cm_end_date")
        if subject_col and cm_name_col and cm_start_col and cm_end_col and all(c in df.columns for c in [subject_col, cm_name_col, cm_start_col, cm_end_col]):
            temp = pd.DataFrame({
                "subject": df[subject_col].astype(str),
                "site": df[site_col] if site_col and site_col in df.columns else "",
                "visit": df[visit_col] if visit_col and visit_col in df.columns else "Screening",
                "drug": df[cm_name_col].astype(str),
                "start": df[cm_start_col].map(to_dt),
                "end": df[cm_end_col].map(to_dt),
            })
            bad = temp[(temp["drug"].astype(str).str.strip() != "") & temp["start"].notna() & temp["end"].notna() & (temp["start"] > temp["end"])]
            for _, r in bad.iterrows():
                add_finding(r["subject"], r["site"], r["visit"], "Concomitant Medication",
                            "CM start date later than end date",
                            "合并用药开始日期晚于结束日期。",
                            sheet_name, f"{cm_start_col} / {cm_end_col}", f"{r['start']} / {r['end']}",
                            "CM Start Date should not be later than End Date", "Not compliant", "Major",
                            "CM temporal order conflict", "No")

    # 跨Sheet主线逻辑：按受试者汇总最早/最晚日期
    def aggregate_subject_field(std_field: str) -> Optional[pd.DataFrame]:
        pairs = get_mapped_infos(mapping_df, std_field)
        frames = []
        for s, c in pairs:
            df = sheet_dict.get(s)
            subj_col = get_local_col(mapping_df, s, "subject_id")
            if df is None or df.empty or not subj_col or subj_col not in df.columns or c not in df.columns:
                continue
            frames.append(pd.DataFrame({"subject": df[subj_col].astype(str), std_field: df[c], "sheet": s}))
        if not frames:
            return None
        return pd.concat(frames, ignore_index=True)

    scr_df = aggregate_subject_field("screening_date")
    rand_df = aggregate_subject_field("randomization_date")
    dose_df = aggregate_subject_field("first_dose_date")
    wd_df = aggregate_subject_field("withdrawal_date")

    if scr_df is not None and rand_df is not None:
        temp = scr_df.merge(rand_df, on="subject", how="inner", suffixes=("_scr", "_rand"))
        temp["screening_date"] = temp["screening_date"].map(to_dt)
        temp["randomization_date"] = temp["randomization_date"].map(to_dt)
        bad = temp[(temp["screening_date"].notna()) & (temp["randomization_date"].notna()) & (temp["randomization_date"] < temp["screening_date"])]
        for _, r in bad.iterrows():
            add_finding(r["subject"], "", "Screening", "Logic",
                        "Randomization Date earlier than Screening Date",
                        "存在关键时间顺序倒挂。",
                        "Cross-Sheet",
                        "screening_date / randomization_date",
                        f"{r['screening_date']} / {r['randomization_date']}",
                        "Screening should occur before Randomization",
                        "Not compliant", "Major", "同一受试者的筛选日期晚于随机日期", "No")

    if dose_df is not None and wd_df is not None:
        temp = dose_df.merge(wd_df, on="subject", how="inner", suffixes=("_dose", "_wd"))
        temp["first_dose_date"] = temp["first_dose_date"].map(to_dt)
        temp["withdrawal_date"] = temp["withdrawal_date"].map(to_dt)
        bad = temp[(temp["first_dose_date"].notna()) & (temp["withdrawal_date"].notna()) & (temp["first_dose_date"] > temp["withdrawal_date"])]
        for _, r in bad.iterrows():
            add_finding(r["subject"], "", "Screening", "Dosing",
                        "Dosing after Withdrawal Date",
                        "给药与退出状态存在冲突。",
                        "Cross-Sheet",
                        "first_dose_date / withdrawal_date",
                        f"{r['first_dose_date']} / {r['withdrawal_date']}",
                        "No dosing should occur after withdrawal",
                        "Not compliant", "Critical", "给药日期晚于退出日期", "No")

    findings_df = pd.DataFrame(findings)
    if findings_df.empty:
        findings_df = pd.DataFrame(columns=["Finding ID", "Subject ID", "Site", "Visit", "Category", "Title", "Description", "Sheet", "Field", "Current Value", "Protocol Requirement", "Judgement", "Risk Level", "Evidence", "Need Medical Review"])

    queries = []
    priority_map = {"Critical": "High", "Major": "High", "Minor": "Medium", "Observation": "Low"}
    for _, f in findings_df.iterrows():
        if f["Risk Level"] in ["Critical", "Major"]:
            queries.append({
                "Subject ID": f["Subject ID"], "Site": f["Site"], "Visit": f["Visit"], "Form/Sheet": f["Sheet"], "Field": f["Field"],
                "Issue": f["Title"],
                "Query Text": f"Please verify the data related to '{f['Title']}' and clarify or update the record if needed based on protocol and source documents.",
                "Query Type": f["Category"], "Suggested Priority": priority_map[f["Risk Level"]],
            })
    queries_df = pd.DataFrame(queries)
    if queries_df.empty:
        queries_df = pd.DataFrame(columns=["Subject ID", "Site", "Visit", "Form/Sheet", "Field", "Issue", "Query Text", "Query Type", "Suggested Priority"])

    subjects = []
    subject_ids = findings_df["Subject ID"].dropna().astype(str).unique().tolist() if not findings_df.empty else []
    for sub in subject_ids:
        sf = findings_df[findings_df["Subject ID"].astype(str) == str(sub)]
        eligibility = min(35, int((sf["Category"] == "Eligibility").sum()) * 20 + int((sf["Category"] == "Logic").sum()) * 10)
        visit_score = min(20, int((sf["Category"] == "Visit").sum()) * 10 + int((sf["Category"] == "Dosing").sum()) * 10)
        safety = min(30, int((sf["Category"] == "Safety").sum()) * 15)
        data_int = min(15, int((sf["Category"].isin(["Data Integrity", "Laboratory", "Concomitant Medication"])).sum()) * 5)
        total = eligibility + visit_score + safety + data_int
        if total >= 80:
            level = "Critical"
        elif total >= 60:
            level = "High"
        elif total >= 40:
            level = "Moderate"
        elif total >= 20:
            level = "Moderate-Low"
        else:
            level = "Low"
        drivers = ", ".join(sf["Title"].head(3).tolist())
        subjects.append({
            "Subject ID": sub, "Eligibility Risk": eligibility, "Visit Risk": visit_score, "Safety Risk": safety,
            "Data Integrity Risk": data_int, "Total Risk Score": total, "Risk Level": level,
            "Top Risk Drivers": drivers, "Recommended Action": "Focused review" if total >= 40 else "Routine follow-up",
        })
    subj_df = pd.DataFrame(subjects)
    if subj_df.empty:
        subj_df = pd.DataFrame(columns=["Subject ID", "Eligibility Risk", "Visit Risk", "Safety Risk", "Data Integrity Risk", "Total Risk Score", "Risk Level", "Top Risk Drivers", "Recommended Action"])

    sites = []
    if not findings_df.empty and "Site" in findings_df.columns:
        for s in findings_df["Site"].fillna("Unmapped").astype(str).unique().tolist():
            sf = findings_df[findings_df["Site"].fillna("Unmapped").astype(str) == s]
            count_sub = len(sf["Subject ID"].fillna("").astype(str).unique())
            mc = int(sf["Risk Level"].isin(["Critical", "Major"]).sum())
            avg_sub = float(subj_df[subj_df["Subject ID"].astype(str).isin(sf["Subject ID"].fillna("").astype(str))]["Total Risk Score"].mean()) if not subj_df.empty else 0
            elig_rate = round(float((sf["Category"] == "Eligibility").mean() * 100), 2) if len(sf) else 0
            safety_rate = round(float((sf["Category"] == "Safety").mean() * 100), 2) if len(sf) else 0
            site_score = round(0.40 * avg_sub + 0.25 * mc + 0.20 * elig_rate + 0.15 * safety_rate, 2)
            if site_score >= 80:
                lvl = "Critical"
            elif site_score >= 60:
                lvl = "High"
            elif site_score >= 40:
                lvl = "Moderate"
            elif site_score >= 20:
                lvl = "Moderate-Low"
            else:
                lvl = "Low"
            sites.append({
                "Site": s, "受试者数": count_sub, "平均风险分": round(avg_sub, 2), "Major/Critical数量": mc,
                "Eligibility高风险比例": elig_rate, "Safety高风险比例": safety_rate,
                "Site Risk Score": site_score, "Site Risk Level": lvl,
            })
    site_df = pd.DataFrame(sites)
    if site_df.empty:
        site_df = pd.DataFrame(columns=["Site", "受试者数", "平均风险分", "Major/Critical数量", "Eligibility高风险比例", "Safety高风险比例", "Site Risk Score", "Site Risk Level"])

    return findings_df, queries_df, subj_df, site_df


def to_excel_bytes(df_dict: Dict[str, pd.DataFrame]) -> bytes:
    bio = io.BytesIO()
    with pd.ExcelWriter(bio, engine='openpyxl') as writer:
        for sheet_name, df in df_dict.items():
            safe_name = re.sub(r"[\\/*?:\[\]]", "_", str(sheet_name))[:31] if sheet_name else 'Sheet1'
            df.to_excel(writer, index=False, sheet_name=safe_name)
    bio.seek(0)
    return bio.getvalue()


st.title("EDC智能稽查与RBQM风险识别系统（MVP）")
st.caption("支持：Excel 全 Sheet 读取、Screening 数据过滤、多 Sheet 实际核查、字段映射、Findings / Query / RBQM / Word 报告")

with st.sidebar:
    st.header("项目创建")
    project_name = st.text_input("项目名称", value="HDM1005-301")
    protocol_code = st.text_input("方案编号", value="HDM1005-301")
    sponsor = st.text_input("申办方", value="")
    analysis_type = st.selectbox("分析类型", ["EDC远程稽查", "数据质量审阅", "安全性专项核查", "入排专项核查"])
    screening_only = st.checkbox("仅分析 Screening / 筛选期 数据", value=True)

col1, col2 = st.columns(2)
with col1:
    protocol_file = st.file_uploader("上传方案（PDF / DOCX）", type=["pdf", "docx"])
with col2:
    edc_file = st.file_uploader("上传 EDC Excel（XLSX / XLS）", type=["xlsx", "xls"])

project_meta = {
    "project_name": project_name,
    "protocol_code": protocol_code,
    "sponsor": sponsor,
    "analysis_type": analysis_type,
    "screening_only": screening_only,
}

raw_sheet_dict: Dict[str, pd.DataFrame] = {}
sheet_dict: Dict[str, pd.DataFrame] = {}
sheet_meta_df = pd.DataFrame(columns=["sheet_name", "row_count", "column_count", "has_data", "inferred_module", "read_status"])
mapping_df = pd.DataFrame(columns=["standard_field_name", "sheet_name", "original_column_name", "mapping_confidence", "score", "is_decode", "needs_review"])
protocol_text = ""
protocol_meta = {"protocol_name": "未识别", "version": "未识别", "version_date": "未识别"}
rules_df = pd.DataFrame()
filtered_sheet_meta_df = pd.DataFrame()

if edc_file is not None:
    try:
        xls = pd.ExcelFile(edc_file)
        metas = []
        for s in xls.sheet_names:
            df = xls.parse(sheet_name=s)
            raw_sheet_dict[s] = df
            metas.append(SheetMeta(
                sheet_name=s,
                row_count=int(df.shape[0]),
                column_count=int(df.shape[1]),
                has_data=bool(df.shape[0] > 0 and df.shape[1] > 0),
                inferred_module=infer_module(s, df.columns.tolist()),
                read_status="成功",
            ).__dict__)
        sheet_meta_df = pd.DataFrame(metas)
        mapping_df = recommend_field_mapping(raw_sheet_dict)
        sheet_dict = build_screening_sheet_dict(raw_sheet_dict, mapping_df) if screening_only else raw_sheet_dict
        filtered_sheet_meta_df = pd.DataFrame([
            {
                "sheet_name": s,
                "row_count": int(df.shape[0]),
                "column_count": int(df.shape[1]),
                "has_data": bool(df.shape[0] > 0 and df.shape[1] > 0),
                "inferred_module": infer_module(s, df.columns.tolist()),
                "read_status": "成功",
            }
            for s, df in sheet_dict.items()
        ])
    except Exception as e:
        st.error(f"Excel 读取失败：{e}")

if protocol_file is not None:
    try:
        protocol_text = read_protocol_text(protocol_file)
        protocol_meta = extract_protocol_meta(protocol_text)
        rules_df = heuristic_rule_extract(protocol_text)
    except Exception as e:
        st.error(f"方案解析失败：{e}")

restricted, reasons = restricted_mode(sheet_meta_df, mapping_df)
if screening_only and filtered_sheet_meta_df is not None and not filtered_sheet_meta_df.empty:
    empty_after_filter = filtered_sheet_meta_df[filtered_sheet_meta_df["row_count"] == 0]["sheet_name"].tolist()
    if empty_after_filter:
        reasons.append("部分 Sheet 在 Screening 过滤后无数据：" + ", ".join(empty_after_filter[:10]))
        restricted = True

page1, page2, page3, page4, page5, page6 = st.tabs([
    "1. 项目与上传", "2. Sheet 读取结果", "3. 字段映射中心", "4. 方案规则中心", "5. 自动核查结果", "6. RBQM 与报告"
])

with page1:
    st.subheader("项目创建页")
    st.json(project_meta)
    st.write({"Analysis Scope": "Screening only" if screening_only else "All visits"})
    if protocol_file:
        st.success(f"已上传方案：{protocol_file.name}")
        st.write(protocol_meta)
    if edc_file:
        st.success(f"已上传 EDC：{edc_file.name}")
        st.write({"实际读取到的 Sheet 数量": len(sheet_meta_df)})
    if restricted:
        st.warning("本次结果为受限分析")
        for r in reasons:
            st.write(f"- {r}")

with page2:
    st.subheader("Excel 全 Sheet 读取与展示")
    if sheet_meta_df.empty:
        st.info("请先上传 EDC Excel")
    else:
        st.markdown("#### 原始读取结果（全部 Sheet）")
        st.dataframe(sheet_meta_df, use_container_width=True)
        if raw_sheet_dict:
            st.download_button(
                "导出全部原始 Sheet（多 Sheet Excel）",
                data=to_excel_bytes(raw_sheet_dict),
                file_name="EDC_All_Sheets_Raw.xlsx",
            )
        if screening_only and filtered_sheet_meta_df is not None and not filtered_sheet_meta_df.empty:
            st.markdown("#### Screening 过滤后结果")
            st.dataframe(filtered_sheet_meta_df, use_container_width=True)
            st.download_button(
                "导出 Screening 数据（多 Sheet Excel）",
                data=to_excel_bytes(sheet_dict),
                file_name="EDC_Screening_Only_All_Sheets.xlsx",
            )
        preview_source = st.radio("预览数据来源", ["原始数据", "Screening 数据" if screening_only else "原始数据"], horizontal=True)
        current_dict = raw_sheet_dict if preview_source == "原始数据" else sheet_dict
        choice = st.selectbox("选择 Sheet 预览", list(current_dict.keys()))
        if choice in current_dict:
            st.write(f"列数：{current_dict[choice].shape[1]} | 行数：{current_dict[choice].shape[0]}")
            st.dataframe(current_dict[choice].head(50), use_container_width=True)

with page3:
    st.subheader("字段映射表")
    if mapping_df.empty:
        st.info("尚未生成字段映射。")
    else:
        edited_mapping = st.data_editor(mapping_df, use_container_width=True, num_rows="dynamic")
        mapping_df = edited_mapping.copy()
        st.download_button("下载字段映射表", data=to_excel_bytes({"field_mapping": mapping_df}), file_name="Field_Mapping.xlsx")

with page4:
    st.subheader("方案规则抽取")
    if not protocol_text:
        st.info("请先上传方案文件。")
    else:
        st.write(protocol_meta)
        st.text_area("方案文本预览（前 5000 字符）", protocol_text[:5000], height=200)
        edited_rules = st.data_editor(rules_df, use_container_width=True, num_rows="dynamic")
        rules_df = edited_rules.copy()
        st.download_button("下载规则表", data=to_excel_bytes({"protocol_rules": rules_df}), file_name="Protocol_Rules.xlsx")

with page5:
    st.subheader("自动核查结果")
    st.caption("当前默认仅对 Screening / 筛选期 数据运行核查。")
    if st.button("运行全流程核查", type="primary"):
        findings_df, queries_df, subj_df, site_df = run_audit(sheet_dict, mapping_df, rules_df)
        st.session_state["findings_df"] = findings_df
        st.session_state["queries_df"] = queries_df
        st.session_state["subj_df"] = subj_df
        st.session_state["site_df"] = site_df

    findings_df = st.session_state.get("findings_df", pd.DataFrame())
    queries_df = st.session_state.get("queries_df", pd.DataFrame())

    coverage_df = build_sheet_coverage(sheet_meta_df, filtered_sheet_meta_df, mapping_df, findings_df if findings_df is not None else pd.DataFrame())
    if coverage_df is not None and not coverage_df.empty:
        st.markdown("#### Sheet Coverage")
        st.dataframe(coverage_df, use_container_width=True)
        st.download_button("导出 Sheet Coverage", data=to_excel_bytes({"Sheet_Coverage": coverage_df}), file_name="Sheet_Coverage.xlsx")

    if findings_df is not None and not findings_df.empty:
        st.markdown("#### Findings")
        st.dataframe(findings_df, use_container_width=True)
        st.download_button("导出 Findings", data=to_excel_bytes({"Findings": findings_df}), file_name="Findings.xlsx")
    else:
        st.info("尚未生成 Findings。")

    if queries_df is not None and not queries_df.empty:
        st.markdown("#### Query")
        st.dataframe(queries_df, use_container_width=True)
        st.download_button("导出 Query", data=to_excel_bytes({"Query": queries_df}), file_name="Query.xlsx")




def build_capa_df(findings_df: pd.DataFrame) -> pd.DataFrame:
    if findings_df is None or findings_df.empty:
        return pd.DataFrame(columns=["CAPA ID","Finding ID","Subject ID","Site","Risk Level","Issue Summary","Root Cause Hypothesis","Corrective Action","Preventive Action","Owner","Target Date","Status"])
    rows=[]
    owner_map={"Critical":"医学监查/数据管理/项目经理","Major":"数据管理/项目经理","Minor":"项目经理/研究中心","Observation":"研究中心"}
    due_map={"Critical":"7天内","Major":"14天内","Minor":"30天内","Observation":"下次例行复核前"}
    for i,(_,r) in enumerate(findings_df.iterrows(),1):
        risk=str(r.get('Risk Level','Observation'))
        category=str(r.get('Category',''))
        title=str(r.get('Title',''))
        descr=str(r.get('Description',''))
        root='需结合源文件、EDC录入路径及研究中心执行流程进一步确认。'
        ca='请核对源文件、方案要求及EDC记录，必要时更正EDC并补充说明。'
        pa='建议对相关岗位进行针对性培训，并增加关键字段复核或系统校验规则。'
        if category == 'Safety':
            root='可能与安全性判定标准理解不一致、随访不完整或字段录入遗漏有关。'
            ca='请核对AE/SAE/AESI原始记录、医学判断及随访状态，必要时补录或更正。'
            pa='建议强化安全性术语培训，并对SAE/AESI关键字段设置复核点。'
        elif category == 'Dosing':
            root='可能与给药主线日期维护不完整、访视执行偏差或录入顺序错误有关。'
            ca='请核对给药记录、药物管理记录及访视记录，必要时修正日期与剂量信息。'
            pa='建议建立给药日期与主线日期的交叉校验机制。'
        elif category == 'Logic':
            root='可能与关键主线字段维护不一致、跨表同步不及时或映射错误有关。'
            ca='请逐一核对相关表单关键日期及受试者状态，统一更正冲突记录。'
            pa='建议增加跨表逻辑检查清单并定期复核。'
        elif category == 'Laboratory':
            root='可能与实验室结果录入、单位维护、异常处理或方案阈值理解偏差有关。'
            ca='请核对原始化验单、单位、参考范围及异常处理记录，必要时补充AE或处理说明。'
            pa='建议对关键实验室指标增加阈值预警与复核流程。'
        rows.append({
            'CAPA ID': f'CAPA-{i:04d}',
            'Finding ID': r.get('Finding ID',''),
            'Subject ID': r.get('Subject ID',''),
            'Site': r.get('Site',''),
            'Risk Level': risk,
            'Issue Summary': title or descr[:80],
            'Root Cause Hypothesis': root,
            'Corrective Action': ca,
            'Preventive Action': pa,
            'Owner': owner_map.get(risk,'项目经理'),
            'Target Date': due_map.get(risk,'30天内'),
            'Status': 'Open',
        })
    return pd.DataFrame(rows)


def build_subject_evidence_summary(findings_df: pd.DataFrame) -> pd.DataFrame:
    if findings_df is None or findings_df.empty:
        return pd.DataFrame(columns=['Subject ID','Site','Findings Count','Highest Risk','Sheets Involved','Categories','Key Evidence Summary'])
    risk_rank={'Critical':4,'Major':3,'Minor':2,'Observation':1}
    rows=[]
    for subject,grp in findings_df.groupby(findings_df['Subject ID'].fillna('').astype(str)):
        if subject == '':
            continue
        highest='Observation'
        best=0
        for x in grp['Risk Level'].fillna('Observation'):
            if risk_rank.get(str(x),0)>best:
                best=risk_rank.get(str(x),0); highest=str(x)
        sheets='；'.join(sorted(set(grp['Sheet'].fillna('').astype(str))))
        cats='；'.join(sorted(set(grp['Category'].fillna('').astype(str))))
        evidences=[]
        for _,r in grp.head(5).iterrows():
            evidences.append(f"{r.get('Title','')}（{r.get('Sheet','')}/{r.get('Field','')}）")
        rows.append({
            'Subject ID': subject,
            'Site': grp['Site'].fillna('').astype(str).iloc[0] if 'Site' in grp.columns and len(grp)>0 else '',
            'Findings Count': len(grp),
            'Highest Risk': highest,
            'Sheets Involved': sheets,
            'Categories': cats,
            'Key Evidence Summary': '；'.join(evidences),
        })
    return pd.DataFrame(rows).sort_values(['Highest Risk','Findings Count'], ascending=[False,False], key=lambda s: s.map(risk_rank).fillna(0) if s.name=='Highest Risk' else s)


def add_df_table(doc, df: pd.DataFrame, title: str, max_rows: int = 200):
    doc.add_paragraph(title)
    if df is None or df.empty:
        doc.add_paragraph('无')
        return
    show=df.head(max_rows).copy()
    t=doc.add_table(rows=1, cols=len(show.columns))
    t.style='Table Grid'
    for i,c in enumerate(show.columns):
        t.rows[0].cells[i].text=str(c)
    for _,r in show.iterrows():
        cells=t.add_row().cells
        for i,c in enumerate(show.columns):
            cells[i].text=str(r.get(c,''))

def gen_word_report(project_meta, protocol_meta, sheet_meta_df, filtered_sheet_meta_df, mapping_df, findings_df, queries_df, subj_df, site_df, restricted, reasons):
    doc = Document()
    try:
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(10.5)
    except Exception:
        pass

    coverage_df = build_sheet_coverage(sheet_meta_df, filtered_sheet_meta_df, mapping_df, findings_df)
    capa_df = build_capa_df(findings_df)
    subject_summary_df = build_subject_evidence_summary(findings_df)

    doc.add_heading('EDC智能稽查报告', level=0)
    p0 = doc.add_paragraph()
    p0.add_run(f"项目名称：{project_meta.get('project_name', '')}\n")
    p0.add_run(f"方案信息：{protocol_meta.get('protocol_name', '')} {protocol_meta.get('version', '')}\n")
    p0.add_run(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}\n")
    p0.add_run(f"分析范围：{'仅Screening/筛选期' if project_meta.get('screening_only', True) else '全部访视'}\n")
    p0.add_run(f"分析模式：{'受限分析' if restricted else '标准分析'}")

    doc.add_heading('1. 项目与稽查概况', level=1)
    t = doc.add_table(rows=1, cols=2)
    t.style = 'Table Grid'
    t.rows[0].cells[0].text='字段'; t.rows[0].cells[1].text='内容'
    rows=[
        ('项目名称', str(project_meta.get('project_name',''))),
        ('项目编号', str(project_meta.get('protocol_code',''))),
        ('申办方', str(project_meta.get('sponsor',''))),
        ('方案名称', str(protocol_meta.get('protocol_name',''))),
        ('方案版本', str(protocol_meta.get('protocol_version',''))),
        ('已读取Sheet数', str(len(sheet_meta_df) if sheet_meta_df is not None else 0)),
    ]
    for k,v in rows:
        row=t.add_row().cells; row[0].text=k; row[1].text=v

    doc.add_heading('2. 分析范围与限制', level=1)
    doc.add_paragraph('本报告基于当前系统实际可读取的方案与EDC数据生成，结论均应结合源文件、研究中心说明及医学判断进一步确认。')
    if restricted:
        doc.add_paragraph('本次为受限分析，主要原因如下：')
        for r in reasons or []:
            doc.add_paragraph(str(r), style='List Bullet')
    else:
        doc.add_paragraph('当前未触发受限分析条件。')

    doc.add_heading('3. Sheet读取与覆盖情况', level=1)
    add_df_table(doc, coverage_df[['Sheet','Rows Read','Rows After Screening','Mapped Fields','Findings Count','Coverage Status']] if coverage_df is not None and not coverage_df.empty else coverage_df, '3.1 Sheet Coverage Summary')

    doc.add_heading('4. 受试者级多Sheet证据链汇总', level=1)
    doc.add_paragraph('以下内容按受试者聚合来自多个Sheet的发现，用于支持跨表证据链审阅。')
    add_df_table(doc, subject_summary_df, '4.1 受试者证据链摘要', max_rows=300)

    doc.add_heading('5. 稽查发现汇总', level=1)
    if findings_df is None or findings_df.empty:
        doc.add_paragraph('未生成发现。')
    else:
        doc.add_paragraph(f"发现总数：{len(findings_df)}")
        if 'Risk Level' in findings_df.columns:
            risk_order=['Critical','Major','Minor','Observation']
            for lvl in risk_order:
                cnt=int((findings_df['Risk Level'].fillna('')==lvl).sum())
                doc.add_paragraph(f"{lvl}: {cnt}", style='List Bullet')
        cols=[c for c in ['Finding ID','Subject ID','Site','Visit','Category','Title','Description','Sheet','Field','Current Value','Protocol Requirement','Judgement','Risk Level','Evidence','Need Medical Review'] if c in findings_df.columns]
        add_df_table(doc, findings_df[cols], '5.1 详细发现', max_rows=500)

    doc.add_heading('6. CAPA建议与整改跟踪', level=1)
    doc.add_paragraph('以下CAPA为系统基于发现自动生成的整改建议，供项目团队、数据管理与医学团队复核后使用。')
    add_df_table(doc, capa_df, '6.1 CAPA建议表', max_rows=500)

    doc.add_heading('7. Query建议', level=1)
    qcols=[c for c in ['Subject ID','Site','Visit','Form/Sheet','Field','Issue','Query Text','Suggested Priority'] if queries_df is not None and c in queries_df.columns]
    add_df_table(doc, queries_df[qcols] if queries_df is not None and not queries_df.empty and qcols else queries_df, '7.1 Query列表', max_rows=300)

    doc.add_heading('8. RBQM风险评分', level=1)
    add_df_table(doc, subj_df, '8.1 受试者级风险评分', max_rows=300)
    add_df_table(doc, site_df, '8.2 中心级风险评分', max_rows=200)

    doc.add_heading('9. 结论与建议', level=1)
    if findings_df is None or findings_df.empty:
        doc.add_paragraph('当前未形成可报告的稽查发现。建议结合源文件及更多EDC模块继续核查。')
    else:
        top_risk='Observation'
        for lvl in ['Critical','Major','Minor','Observation']:
            if (findings_df['Risk Level'].fillna('')==lvl).any():
                top_risk=lvl; break
        doc.add_paragraph(f'本次稽查最高风险等级为：{top_risk}。')
        doc.add_paragraph('建议优先处理Critical/Major问题，并对涉及安全性、给药主线、资格判断及关键时间逻辑的事项开展专项复核。')
        doc.add_paragraph('对于需要医学复核的问题，应由医学监查或具备资质的专业人员最终确认。')

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.getvalue()

with page6:
    st.subheader("RBQM评分、CAPA与Word报告")
    findings_df = st.session_state.get("findings_df", pd.DataFrame())
    queries_df = st.session_state.get("queries_df", pd.DataFrame())
    subj_df = st.session_state.get("subj_df", pd.DataFrame())
    site_df = st.session_state.get("site_df", pd.DataFrame())
    capa_df = build_capa_df(findings_df)
    subject_summary_df = build_subject_evidence_summary(findings_df)

    st.markdown("#### 受试者级多Sheet证据链汇总")
    st.dataframe(subject_summary_df, use_container_width=True)
    if subject_summary_df is not None and not subject_summary_df.empty:
        st.download_button("导出受试者证据链汇总", data=to_excel_bytes({"Subject_Evidence_Summary": subject_summary_df}), file_name="Subject_Evidence_Summary.xlsx")

    c1, c2 = st.columns(2)
    with c1:
        st.markdown("#### 受试者风险评分")
        st.dataframe(subj_df, use_container_width=True)
    with c2:
        st.markdown("#### 中心风险评分")
        st.dataframe(site_df, use_container_width=True)

    st.markdown("#### CAPA建议")
    st.dataframe(capa_df, use_container_width=True)
    if capa_df is not None and not capa_df.empty:
        st.download_button("导出CAPA建议", data=to_excel_bytes({"CAPA": capa_df}), file_name="CAPA_Plan.xlsx")

    if findings_df is not None and not findings_df.empty:
        report_bytes = gen_word_report(project_meta, protocol_meta, sheet_meta_df, filtered_sheet_meta_df if screening_only else sheet_meta_df, mapping_df, findings_df, queries_df, subj_df, site_df, restricted, reasons)
        st.download_button("下载中文版Word稽查报告", data=report_bytes, file_name="EDC智能稽查报告_中文版.docx")

st.markdown("---")
st.caption("说明：当前版本已支持多Sheet实际核查、按受试者聚合证据链、自动生成CAPA建议及更接近正式稽查报告的中文版Word导出。")
